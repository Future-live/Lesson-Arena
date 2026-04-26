import html
import logging
import shutil
import socket
import subprocess
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import bleach
import markdown
from django.conf import settings
from docx import Document as DocxDocument
from pypdf import PdfReader

from apps.lesson_plans.models import LessonPlanDocument


logger = logging.getLogger(__name__)

ALLOWED_HTML_TAGS = set(bleach.sanitizer.ALLOWED_TAGS) | {
    "p",
    "h1",
    "h2",
    "h3",
    "h4",
    "section",
    "pre",
    "br",
    "table",
    "tbody",
    "tr",
    "th",
    "td",
    "ul",
    "ol",
    "li",
    "strong",
    "em",
}
ALLOWED_HTML_ATTRIBUTES = {
    "a": ["href", "title", "target", "rel"],
}


@dataclass
class ParseResult:
    display_mode: str
    extracted_text: str
    rendered_html: str
    word_count: int
    page_count: int
    preview_pdf_bytes: bytes | None = None


def sanitize_html(content: str) -> str:
    return bleach.clean(
        content,
        tags=ALLOWED_HTML_TAGS,
        attributes=ALLOWED_HTML_ATTRIBUTES,
        strip=True,
    )


def count_non_whitespace(text: str) -> int:
    return sum(1 for char in text if not char.isspace())


def paragraphs_to_html(paragraphs: list[str]) -> str:
    blocks = [f"<p>{html.escape(paragraph)}</p>" for paragraph in paragraphs if paragraph]
    return sanitize_html("\n".join(blocks) or "<p>文档未抽取到可展示内容。</p>")


def count_pdf_pages(pdf_bytes: bytes) -> int:
    return len(PdfReader(BytesIO(pdf_bytes)).pages)


def convert_office_document_to_pdf(document: LessonPlanDocument) -> bytes:
    office_binary = shutil.which("soffice") or shutil.which("libreoffice")
    if not office_binary:
        raise ValueError("服务器未安装 LibreOffice，当前无法生成 Word 版式预览。")

    suffix = Path(document.original_filename).suffix.lower()
    with tempfile.TemporaryDirectory() as temp_dir:
        source_path = Path(temp_dir) / f"source{suffix}"
        with source_path.open("wb") as file_obj:
            for chunk in document.original_file.chunks():
                file_obj.write(chunk)
        document.original_file.seek(0)

        result = subprocess.run(
            [
                office_binary,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                temp_dir,
                str(source_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )

        preview_path = source_path.with_suffix(".pdf")
        if result.returncode != 0 or not preview_path.exists():
            raise ValueError(result.stderr.strip() or "Word 文档转 PDF 预览失败。")

        return preview_path.read_bytes()


def parse_pdf(document: LessonPlanDocument) -> ParseResult:
    reader = PdfReader(document.original_file)
    document.original_file.seek(0)
    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.PDF,
        extracted_text="",
        rendered_html="",
        word_count=0,
        page_count=len(reader.pages),
    )


def parse_docx(document: LessonPlanDocument) -> ParseResult:
    try:
        preview_pdf_bytes = convert_office_document_to_pdf(document)
    except ValueError as exc:
        if "LibreOffice" not in str(exc):
            raise
        logger.warning(
            "LibreOffice 不可用，降级使用 python-docx 渲染 HTML 预览。document_id=%s",
            document.id,
        )
        return parse_docx_as_html(document)

    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.PDF,
        extracted_text="",
        rendered_html="",
        word_count=0,
        page_count=count_pdf_pages(preview_pdf_bytes),
        preview_pdf_bytes=preview_pdf_bytes,
    )


def parse_docx_as_html(document: LessonPlanDocument) -> ParseResult:
    raw_bytes = document.original_file.read()
    document.original_file.seek(0)
    docx = DocxDocument(BytesIO(raw_bytes))
    text_parts: list[str] = []
    html_blocks: list[str] = []

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        text_parts.append(text)
        html_blocks.append(f"<p>{html.escape(text)}</p>")

    for table in docx.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(part.text.strip() for part in cell.paragraphs if part.text.strip())
                if cell_text:
                    text_parts.append(cell_text)
                cells.append(f"<td>{html.escape(cell_text)}</td>")
            rows.append(f"<tr>{''.join(cells)}</tr>")
        if rows:
            html_blocks.append(f"<table><tbody>{''.join(rows)}</tbody></table>")

    extracted_text = "\n".join(text_parts)
    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.HTML,
        extracted_text=extracted_text,
        rendered_html=sanitize_html("\n".join(html_blocks) or "<p>文档未抽取到可展示内容。</p>"),
        word_count=count_non_whitespace(extracted_text),
        page_count=1 if extracted_text else 0,
    )


def parse_legacy_doc(document: LessonPlanDocument) -> ParseResult:
    preview_pdf_bytes = convert_office_document_to_pdf(document)
    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.PDF,
        extracted_text="",
        rendered_html="",
        word_count=0,
        page_count=count_pdf_pages(preview_pdf_bytes),
        preview_pdf_bytes=preview_pdf_bytes,
    )


def parse_text_file(document: LessonPlanDocument) -> ParseResult:
    raw_bytes = document.original_file.read()
    document.original_file.seek(0)
    decoded = None
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            decoded = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    decoded = decoded or raw_bytes.decode("utf-8", errors="ignore")

    paragraphs = [line.strip() for line in decoded.splitlines() if line.strip()]
    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.HTML,
        extracted_text=decoded,
        rendered_html=paragraphs_to_html(paragraphs),
        word_count=count_non_whitespace(decoded),
        page_count=1 if decoded else 0,
    )


def parse_markdown_file(document: LessonPlanDocument) -> ParseResult:
    raw_bytes = document.original_file.read()
    document.original_file.seek(0)
    decoded = raw_bytes.decode("utf-8", errors="ignore")
    rendered_html = markdown.markdown(
        decoded,
        extensions=["fenced_code", "tables", "sane_lists", "toc"],
    )
    return ParseResult(
        display_mode=LessonPlanDocument.DisplayMode.HTML,
        extracted_text=decoded,
        rendered_html=sanitize_html(rendered_html),
        word_count=count_non_whitespace(decoded),
        page_count=1 if decoded else 0,
    )


def parse_document(document: LessonPlanDocument) -> ParseResult:
    extension = Path(document.original_filename).suffix.lower()
    if extension == ".pdf":
        return parse_pdf(document)
    if extension == ".doc":
        return parse_legacy_doc(document)
    if extension == ".docx":
        return parse_docx(document)
    if extension in {".txt"}:
        return parse_text_file(document)
    if extension in {".md", ".markdown"}:
        return parse_markdown_file(document)
    raise ValueError(f"不支持的文件格式：{extension}")


def redis_url_is_available(url: str) -> bool:
    parsed = urlparse(url)
    if parsed.scheme not in {"redis", "rediss"}:
        return True

    host = parsed.hostname or "127.0.0.1"
    port = parsed.port or 6379
    try:
        with socket.create_connection((host, port), timeout=0.25):
            return True
    except OSError:
        return False


def should_parse_synchronously() -> bool:
    if settings.CELERY_TASK_ALWAYS_EAGER:
        return True

    broker_available = redis_url_is_available(settings.CELERY_BROKER_URL)
    backend_available = redis_url_is_available(settings.CELERY_RESULT_BACKEND)
    return not broker_available or not backend_available


def dispatch_document_parse(document_id: str) -> None:
    from apps.lesson_plans.tasks import parse_document_safely, parse_document_task

    if should_parse_synchronously():
        logger.warning("异步队列不可用或配置为同步执行，直接解析文档。document_id=%s", document_id)
        parse_document_safely(document_id)
        return

    try:
        parse_document_task.delay(document_id)
    except Exception:
        logger.exception("异步队列不可用，切换为同步解析。document_id=%s", document_id)
        parse_document_safely(document_id)
